"""Setup pipeline: build challenges.csv and expectations.csv.

Creates output dirs, splits discovery notes, extracts worksheets (one record per
line), merges one-word meeting-note lines into the next note (notes only),
adds cleaned text for sentiment, then writes (nothing is dropped):

  output/raw/worksheets.csv
  output/processed/challenges.csv    (focus_group, pain_points, processed_text)
  output/processed/expectations.csv  (focus_group, expectations, processed_text)

Usage:
  python3 setup/extract_records.py
"""

from __future__ import annotations

import re
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
NOTES_PATH = Path("./data/notes.docx")
WORKSHEETS_DIR = Path("./data/worksheets")
OUTPUT_DIR = Path("./output")
RAW_DIR = OUTPUT_DIR / "raw"
PROCESSED_DIR = OUTPUT_DIR / "processed"
SECTIONS_DIR = RAW_DIR / "docx_sections"
WORKSHEETS_CSV = RAW_DIR / "worksheets.csv"
CHALLENGES_CSV = PROCESSED_DIR  / "challenges.csv"
EXPECTATIONS_CSV = PROCESSED_DIR  / "expectations.csv"

WORKSHEET_FOCUS_GROUPS = {
    "baa supervisors focus group": "BAA Supervisors & Admin Aide",
    "building inspector - round 1 doce": "DOCE Building Inspectors",
    "cpc focus group": "CPC",
    "cpo central permit office": "DOCE Central Permit Office",
    "cpo co-ordinator focus group": "DOCE CPO Coordinators",
    "cpo coordinator focus group": "DOCE CPO Coordinators",
    "doce admin aide": "DOCE Admin Aides",
    "fire department focus group": "DOCE Fire Prevention Bureau",
    "housing inspectors - round 1 doce": "DOCE Housing Inspectors",
    "law": "Law",
    "nbd data team worksheet response": "NBD Data Team",
    "nbd data team": "NBD Data Team",
    "nbd focus group": "NBD Neighborhood Planning",
    "office manager - round 1 doce": "DOCE Office Staff",
    "perm_com_elec inspectors - round 1 doce": "DOCE CommercialPermitElectrical Inspectors",
    "round 1 discover worksheet response zoning": "DOCE Zoning",
    "supervisors - round 1 doce": "DOCE Supervisors",
}

LIKES_HEADER = "what features do you like about ips"
TOOLS_HEADER = "are there other tools or systems"
EXPECTATIONS_HEADER = "what features or capabilities do you wish ips had"

PAIN_HEADING = re.compile(
    r"^(?:pain\s*points?\s*(?:and|&|/)\s*)?challenges?\s*:?\s*$", re.I
)
EXPECTATION_HEADING = re.compile(r"^future\s*capabilities?\s*:?\s*$", re.I)

MIN_WORDS = 4


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------
def ensure_dirs() -> None:
    for path in (OUTPUT_DIR, RAW_DIR, PROCESSED_DIR, SECTIONS_DIR):
        path.mkdir(parents=True, exist_ok=True)


def normalize_text(text: str) -> str:
    return text.strip().lower().replace("\u2019", "'").replace("\xa0", " ")


def normalize_focus_group(name) -> str:
    return re.sub(r"\s+", " ", str(name).replace("\xa0", " ")).strip()


def split_lines(text) -> list[str]:
    if pd.isna(text) or not str(text).strip():
        return []
    return [ln.strip() for ln in str(text).replace("\r\n", "\n").split("\n") if ln.strip()]


def unique_cell_texts(row) -> list[str]:
    """One text per logical Word cell (python-docx repeats merged cells)."""
    texts, seen = [], set()
    for cell in row.cells:
        key = id(cell._tc)
        if key in seen:
            continue
        seen.add(key)
        texts.append(cell.text.replace("\xa0", " ").strip())
    return texts


# ---------------------------------------------------------------------------
# 1. Split discovery notes → output/raw/docx_sections/
# ---------------------------------------------------------------------------
def simple_focus_group_name(heading_text: str, fallback: str = "section") -> str:
    text = str(heading_text).replace("\xa0", " ")
    text = re.sub(r"(?i)\bdiscovery\s+session\b", "", text)
    parts = [p.strip() for p in re.split(r"\s*\|\s*", text) if p.strip()]

    cleaned = []
    for part in parts:
        part = part.replace("/", "")
        part = re.sub(r"(?i)\s+group\s*\d*$", "", part).strip()
        part = re.sub(r"\s+\d+$", "", part).strip()
        part = re.sub(r"\s{2,}", " ", part).strip(" .-_")
        if part:
            cleaned.append(part)

    name = " ".join(cleaned)
    name = re.sub(r"\b(\w+)\s+\1\b", r"\1", name, flags=re.I)
    name = re.sub(r"\s{2,}", " ", name).strip(" .-_")
    name = "".join(ch for ch in name if ch not in '<>:"/\\|?*').strip()
    name = re.sub(r"\s{2,}", " ", name).strip(" .-_")
    return name or fallback


def _copy_paragraph(src_para, dst_doc):
    new_para = dst_doc.add_paragraph()
    try:
        if src_para.style is not None:
            new_para.style = src_para.style.name
    except Exception:
        pass
    for run in src_para.runs:
        new_run = new_para.add_run(run.text)
        for attr in ("bold", "italic", "underline"):
            try:
                setattr(new_run, attr, getattr(run, attr))
            except Exception:
                pass
        try:
            if run.font is not None:
                if run.font.name:
                    new_run.font.name = run.font.name
                if run.font.size:
                    new_run.font.size = run.font.size
        except Exception:
            pass
    return new_para


def split_docx_by_heading(
    input_path: Path | str = NOTES_PATH,
    output_dir: Path | str = SECTIONS_DIR,
) -> dict[str, int]:
    input_path, output_dir = Path(input_path), Path(output_dir)
    if not input_path.exists():
        raise FileNotFoundError(f"Missing notes file: {input_path}")

    doc = Document(input_path)
    sections: list[tuple] = []
    current = None

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading 3"):
            if current is not None:
                sections.append(current)
            current = (para, [])
        elif current is not None:
            current[1].append(para)
    if current is not None:
        sections.append(current)

    output_dir.mkdir(parents=True, exist_ok=True)
    for existing in output_dir.glob("*.docx"):
        existing.unlink()

    written: dict[str, int] = {}
    for index, (heading_para, content_paras) in enumerate(sections, start=1):
        heading_text = heading_para.text if heading_para is not None else f"section_{index}"
        focus_group = simple_focus_group_name(heading_text, fallback=f"section_{index}")
        out_path = output_dir / f"{focus_group}.docx"

        section_doc = Document(out_path) if focus_group in written else Document()
        _copy_paragraph(heading_para, section_doc)
        for paragraph in content_paras:
            _copy_paragraph(paragraph, section_doc)
        section_doc.save(out_path)

        written[focus_group] = written.get(focus_group, 0) + 1
        action = "Merged" if written[focus_group] > 1 else "Saved"
        print(f"  {action}: {out_path.name}")

    print(f"Split {len(written)} focus-group files → {output_dir}")
    return written


# ---------------------------------------------------------------------------
# 2. Extract worksheets → long-form rows
# ---------------------------------------------------------------------------
def focus_group_name_from_filename(filename: str) -> str:
    stem = Path(filename).stem.replace("\xa0", " ")
    stem = re.sub(r"\s+Worksheet(?:\s+Response)?$", "", stem, flags=re.I)
    stem = re.sub(r"\s{2,}", " ", stem).strip()
    key = stem.lower()
    if key in WORKSHEET_FOCUS_GROUPS:
        return WORKSHEET_FOCUS_GROUPS[key]

    cleaned = re.sub(r"(?i)\bfocus\s+group\b", "", stem)
    cleaned = re.sub(r"(?i)\bround\s*1\b", "", cleaned)
    cleaned = re.sub(r"(?i)\bdiscover(?:y)?\b", "", cleaned)
    cleaned = re.sub(r"(?i)\bresponse\b", "", cleaned)
    cleaned = re.sub(r"\s*[-–—]\s*", " ", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip(" .-_")
    return cleaned or stem


def get_table_type(table) -> str | None:
    if not table.rows:
        return None
    cells = unique_cell_texts(table.rows[0])
    first = normalize_text(cells[0]) if cells else ""
    for key in ("role", "process", "pain", "technology"):
        if first.startswith(key):
            return "pain_point" if key == "pain" else key
    return None


def extract_pain_expectation_records(
    table, focus_group: str
) -> tuple[list[dict], list[dict]]:
    """Pain table: left = pain, right = suggested improvements → expectations."""
    pains, expectations = [], []

    for row in table.rows:
        cells = unique_cell_texts(row)
        if not any(cells):
            continue

        joined = normalize_text(" ".join(cells))
        if joined.startswith("pain points") or joined.startswith("pain point"):
            continue
        if "issues or frustrations" in joined and "suggested improvements" in joined:
            continue
        if normalize_text(cells[0]) in {"issues or frustrations", "suggested improvements"}:
            continue

        left = cells[0] if cells else ""
        right = cells[1] if len(cells) > 1 else ""

        for line in split_lines(left):
            pains.append(
                {"focus_group_name": focus_group, "pain_point": line, "expectations": pd.NA}
            )
        for line in split_lines(right):
            expectations.append(
                {"focus_group_name": focus_group, "pain_point": pd.NA, "expectations": line}
            )

    return pains, expectations


def extract_technology_expectation_records(table, focus_group: str) -> list[dict]:
    expectations, section = [], None

    for row in table.rows:
        cells = unique_cell_texts(row)
        if not any(cells):
            continue

        row_text = normalize_text(" ".join(cells))
        first_cell = normalize_text(cells[0])

        if first_cell.startswith("technology"):
            continue
        if LIKES_HEADER in row_text:
            section = "likes_dislikes"
            continue
        if TOOLS_HEADER in row_text or (first_cell == "tool" and "used for" in row_text):
            section = "tools"
            continue
        if EXPECTATIONS_HEADER in row_text:
            section = "expectations"
            continue

        if section != "expectations":
            continue

        for cell_text in cells:
            for line in split_lines(cell_text):
                expectations.append(
                    {
                        "focus_group_name": focus_group,
                        "pain_point": pd.NA,
                        "expectations": line,
                    }
                )

    return expectations


def extract_worksheet_records(doc_path: Path) -> list[dict]:
    if doc_path.name.startswith("~$"):
        return []

    focus_group = focus_group_name_from_filename(doc_path.name)
    records: list[dict] = []

    for table in Document(doc_path).tables:
        table_type = get_table_type(table)
        if table_type == "pain_point":
            pains, exps = extract_pain_expectation_records(table, focus_group)
            records.extend(pains)
            records.extend(exps)
        elif table_type == "technology":
            records.extend(extract_technology_expectation_records(table, focus_group))

    return records


def build_worksheets_df(worksheets_dir: Path | str = WORKSHEETS_DIR) -> pd.DataFrame:
    all_records: list[dict] = []

    for doc_path in sorted(Path(worksheets_dir).rglob("*.docx")):
        if doc_path.name.startswith("~$"):
            continue
        try:
            rows = extract_worksheet_records(doc_path)
            all_records.extend(rows)
            n_pain = sum(pd.notna(r["pain_point"]) for r in rows)
            n_exp = sum(pd.notna(r["expectations"]) for r in rows)
            print(f"  {doc_path.name}: {n_pain} pain + {n_exp} expectation")
        except zipfile.BadZipFile:
            print(f"  Skipping {doc_path.name}: not a valid docx")
        except Exception as exc:
            print(f"  Error reading {doc_path.name}: {exc}")

    return pd.DataFrame(
        all_records, columns=["focus_group_name", "pain_point", "expectations"]
    )


# ---------------------------------------------------------------------------
# 3. Discovery notes → pain / expectation tuples
# ---------------------------------------------------------------------------
def normalize_heading_text(text: str) -> str:
    text = re.sub(r"\s+", " ", str(text).replace("\xa0", " ")).strip().rstrip(":").strip()
    text = re.sub(r"(?i)(challenges)\1+", r"\1", text)
    text = re.sub(r"(?i)(capabilities)bilities", r"\1", text)
    return text


def heading4_section_type(text: str) -> str | None:
    normalized = normalize_heading_text(text)
    if PAIN_HEADING.match(normalized):
        return "pain"
    if EXPECTATION_HEADING.match(normalized):
        return "expectation"
    return None


def records_from_worksheets(worksheets) -> tuple[list[tuple], list[tuple]]:
    df = worksheets if isinstance(worksheets, pd.DataFrame) else pd.read_csv(worksheets)
    pain_points, expectations = [], []

    for row in df.itertuples(index=False):
        focus_group = normalize_focus_group(row.focus_group_name)
        for line in split_lines(getattr(row, "pain_point", None)):
            pain_points.append((focus_group, line))
        for line in split_lines(getattr(row, "expectations", None)):
            expectations.append((focus_group, line))

    print(f"Worksheets: {len(pain_points)} pain · {len(expectations)} expectations")
    return pain_points, expectations


def records_from_discovery_notes(
    discovery_notes_path: Path | str = SECTIONS_DIR,
) -> tuple[list[tuple], list[tuple]]:
    """Extract Heading 4 pain/expectation lines, then merge one-word speaker labels.

    One-word consolidation happens here (meeting-note order) so context stays with
    the following line. Worksheet rows are left alone — cells are independent.
    """
    pain_points, expectations = [], []
    section_files = sorted(Path(discovery_notes_path).glob("*.docx"))

    if not section_files:
        print(f"Discovery notes: 0 files in {discovery_notes_path}")
        return pain_points, expectations

    for doc_path in section_files:
        focus_group = normalize_focus_group(doc_path.stem)
        section = None
        for para in Document(doc_path).paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if style.startswith("Heading"):
                section = (
                    heading4_section_type(text) if style.startswith("Heading 4") else None
                )
                continue
            if not text or section is None:
                continue
            target = pain_points if section == "pain" else expectations
            for line in split_lines(text):
                target.append((focus_group, line))

    before = (len(pain_points), len(expectations))
    if pain_points:
        pain_df = pd.DataFrame(pain_points, columns=["focus_group", "pain_points"])
        pain_df = consolidate_one_word_with_next(pain_df, "pain_points")
        pain_points = list(pain_df.itertuples(index=False, name=None))
    if expectations:
        exp_df = pd.DataFrame(expectations, columns=["focus_group", "expectations"])
        exp_df = consolidate_one_word_with_next(exp_df, "expectations")
        expectations = list(exp_df.itertuples(index=False, name=None))

    print(
        f"Discovery notes: pain {before[0]} → {len(pain_points)}, "
        f"expectations {before[1]} → {len(expectations)} "
        f"(one-word lines merged into following note)"
    )
    return pain_points, expectations


# ---------------------------------------------------------------------------
# 5. Text preprocessing (clean text only — nothing dropped in setup)
# ---------------------------------------------------------------------------
def filter_short(
    df: pd.DataFrame, text_col: str, min_words: int = MIN_WORDS
) -> pd.DataFrame:
    """Available for analysis; setup does not drop rows."""
    return df[df[text_col].astype(str).str.split().str.len() > min_words].reset_index(
        drop=True
    )


def consolidate_one_word_with_next(
    df: pd.DataFrame, text_col: str, group_col: str = "focus_group"
) -> pd.DataFrame:
    """Merge one-word rows into the next multi-word row within the same focus group.

    Intended for discovery-session notes (paragraph order). Do not apply to
    worksheet rows — those cells are separate answers.
    """
    rows = df.reset_index(drop=True).to_dict("records")
    result, i = [], 0

    while i < len(rows):
        group = rows[i][group_col]
        text = "" if pd.isna(rows[i][text_col]) else str(rows[i][text_col]).strip()
        if not text:
            i += 1
            continue

        if len(text.split()) == 1:
            parts = [text]
            j = i + 1
            while j < len(rows) and rows[j][group_col] == group:
                nxt = (
                    ""
                    if pd.isna(rows[j][text_col])
                    else str(rows[j][text_col]).strip()
                )
                if not nxt:
                    j += 1
                    continue
                if len(nxt.split()) == 1:
                    parts.append(nxt)
                    j += 1
                    continue
                rows[j][text_col] = f"{' '.join(parts)} {nxt}".strip()
                i = j
                break
            else:
                result.append({group_col: group, text_col: " ".join(parts)})
                i = j
            continue

        result.append({group_col: group, text_col: text})
        i += 1

    return pd.DataFrame(result, columns=[group_col, text_col])


def clean_for_sentiment(series: pd.Series) -> pd.Series:
    """Normalize whitespace only — keep punctuation, digits, and other characters."""
    return (
        series.fillna("")
        .astype(str)
        .str.replace(r"\s+", " ", regex=True)
        .str.strip()
    )


def sentiment_summary(df: pd.DataFrame, text_col: str) -> pd.DataFrame:
    return (
        df.groupby("focus_group", dropna=False)
        .agg(
            total=(text_col, "count"),
            negative=("sentiment", lambda s: (s == "negative").sum()),
            positive=("sentiment", lambda s: (s == "positive").sum()),
            neutral=("sentiment", lambda s: (s == "neutral").sum()),
        )
        .reset_index()
        .sort_values(["total", "focus_group"], ascending=[False, True])
    )


def preprocess_records(
    challenges_df: pd.DataFrame,
    expectations_df: pd.DataFrame,
) -> tuple[pd.DataFrame, pd.DataFrame]:
    """Add processed_text only. Setup keeps every row (no short-row drop)."""
    challenges_df = challenges_df.copy()
    expectations_df = expectations_df.copy()
    challenges_df["processed_text"] = clean_for_sentiment(challenges_df["pain_points"])
    expectations_df["processed_text"] = clean_for_sentiment(
        expectations_df["expectations"]
    )

    print(
        f"Preprocessed text columns added: "
        f"pain {len(challenges_df)}, exp {len(expectations_df)} (none dropped)"
    )
    return challenges_df, expectations_df


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
def run(
    notes_path: Path | str = NOTES_PATH,
    worksheets_dir: Path | str = WORKSHEETS_DIR,
    sections_dir: Path | str = SECTIONS_DIR,
    worksheets_csv: Path | str = WORKSHEETS_CSV,
    challenges_csv: Path | str = CHALLENGES_CSV,
    expectations_csv: Path | str = EXPECTATIONS_CSV,
    split_notes: bool = True,
    preprocess: bool = True,
) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    """Full setup: dirs → split notes → worksheets → challenges/expectations."""
    ensure_dirs()
    sections_dir = Path(sections_dir)
    worksheets_csv = Path(worksheets_csv)
    challenges_csv = Path(challenges_csv)
    expectations_csv = Path(expectations_csv)

    print("=== 1. Directories ===")
    print(f"  RAW_DIR       → {RAW_DIR.resolve()}")
    print(f"  PROCESSED_DIR → {PROCESSED_DIR.resolve()}")
    print(f"  SECTIONS_DIR  → {sections_dir.resolve()}")

    if split_notes:
        print("\n=== 2. Split discovery notes ===")
        if Path(notes_path).exists():
            split_docx_by_heading(notes_path, sections_dir)
        else:
            print(f"  Skipping: missing {notes_path}")

    print("\n=== 3. Extract worksheets ===")
    worksheets_df = build_worksheets_df(worksheets_dir)
    worksheets_df.to_csv(worksheets_csv, index=False, encoding="utf-8-sig")
    print(f"  Saved {len(worksheets_df)} rows → {worksheets_csv}")

    print("\n=== 4. Build challenges & expectations ===")
    ws_pain, ws_exp = records_from_worksheets(worksheets_df)
    dn_pain, dn_exp = records_from_discovery_notes(sections_dir)

    challenges_df = pd.DataFrame(
        ws_pain + dn_pain, columns=["focus_group", "pain_points"]
    )
    expectations_df = pd.DataFrame(
        ws_exp + dn_exp, columns=["focus_group", "expectations"]
    )

    if preprocess:
        print("\n=== 5. Preprocess (clean text only — nothing dropped) ===")
        challenges_df, expectations_df = preprocess_records(
            challenges_df, expectations_df
        )

    Path(challenges_csv).parent.mkdir(parents=True, exist_ok=True)
    Path(expectations_csv).parent.mkdir(parents=True, exist_ok=True)
    challenges_df.to_csv(challenges_csv, index=False, encoding="utf-8-sig")
    expectations_df.to_csv(expectations_csv, index=False, encoding="utf-8-sig")

    print(f"\nSaved {len(challenges_df)} challenges → {challenges_csv}")
    print(f"Saved {len(expectations_df)} expectations → {expectations_csv}")
    return worksheets_df, challenges_df, expectations_df


if __name__ == "__main__":
    run()
